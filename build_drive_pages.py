"""Parse every still-unused docx in the Data/ drive into proper dashboard pages.

Produces:
  - chiefs.html         — biographical timelines for Crozier + Lewis
  - reports.html        — annual report excerpts (the Board's own voice)
  - tech-master.html    — canonical year-period × technology mapping
  - reading-room.html   — index linking the above

All four are static, paper-themed, brand-matched to the rest of the dashboard.
"""
from __future__ import annotations

import json
import re
from html import escape as h
from pathlib import Path

from docx import Document

ROOT = Path(__file__).parent
DATA = ROOT / "Data"

# ── Brand tokens (mirror the rest of the dashboard) ─────────────────────────
PAPER       = "#F6F6F7"
PAPER_SOFT  = "#FFFFFF"
BORDER      = "#E1E3E6"
TEXT        = "#1A1F36"
TEXT_MID    = "#5A6075"
TEXT_SOFT   = "#8C92A4"
BRASS       = "#C9A24C"
ACCENT      = "#635BFF"


# ── Page shell ──────────────────────────────────────────────────────────────

_SHELL = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} · Fortify the Ordnance</title>
<link rel="icon" type="image/svg+xml" href="data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'><path fill-rule='evenodd' clip-rule='evenodd' d='M32 3 L42 22 L61 32 L42 42 L32 61 L22 42 L3 32 L22 22 Z M32 26 A6 6 0 1 0 32 38 A6 6 0 1 0 32 26 Z' fill='%23C9A24C'/><circle cx='32' cy='32' r='2.2' fill='%23C9A24C'/></svg>">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500;600;700&family=Source+Serif+4:opsz,wght@8..60,400;8..60,600;8..60,700&display=swap" rel="stylesheet">
<style>
  :root {{
    --bg: #F6F6F7;
    --panel: #FFFFFF;
    --panel-2: #FAFAFB;
    --border: #E1E3E6;
    --text: #1A1F36;
    --text-mid: #5A6075;
    --text-soft: #8C92A4;
    --brass: #C9A24C;
    --accent: #635BFF;
    --serif: 'Source Serif 4', Georgia, serif;
    --sans: 'Inter', -apple-system, sans-serif;
    --mono: 'JetBrains Mono', monospace;
  }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  html, body {{
    background: var(--bg);
    color: var(--text);
    font-family: var(--sans);
    font-size: 15.5px;
    line-height: 1.65;
    -webkit-font-smoothing: antialiased;
  }}
  a {{ color: var(--accent); text-decoration: none; }}
  a:hover {{ text-decoration: underline; }}

  .topbar {{
    display: flex; align-items: center; justify-content: space-between;
    padding: 12px 28px;
    background: var(--panel);
    border-bottom: 1px solid var(--border);
    font-family: var(--mono);
    font-size: 11px; letter-spacing: 1.2px; color: var(--text-mid);
  }}
  .tb-l {{
    display: flex; align-items: center; gap: 10px;
    color: inherit; text-decoration: none; cursor: pointer;
  }}
  .tb-l svg {{
    width: 18px; height: 18px; color: var(--brass);
    transition: transform .5s cubic-bezier(.2,.8,.2,1);
  }}
  .tb-l:hover svg {{ transform: rotate(45deg); }}
  .tb-l strong {{ color: var(--text); font-weight: 700; letter-spacing: 1.4px; }}
  .tb-r a {{ color: var(--text-mid); letter-spacing: 1.2px; }}
  .tb-r a:hover {{ color: var(--brass); text-decoration: none; }}

  .wrap {{
    max-width: 1320px;
    margin: 0 auto;
    padding: 56px 48px 96px;
  }}
  .eyebrow {{
    font-family: var(--mono);
    font-size: 10.5px; font-weight: 700; letter-spacing: 2.4px;
    text-transform: uppercase; color: var(--brass);
    margin-bottom: 18px; display: flex; align-items: center; gap: 14px;
  }}
  .eyebrow::before {{
    content: ""; width: 32px; height: 1px; background: var(--brass);
  }}
  h1.page-title {{
    font-family: var(--serif);
    font-weight: 700; font-size: 56px; letter-spacing: -1.6px;
    line-height: 1.04; margin-bottom: 18px; max-width: 1100px;
  }}
  h1.page-title em {{ font-style: italic; color: var(--brass); font-weight: 400; }}
  .lede {{
    font-family: var(--serif);
    font-size: 20px; line-height: 1.5; color: var(--text-mid);
    max-width: 980px; margin-bottom: 12px;
  }}
  footer {{
    margin-top: 88px; padding-top: 22px;
    border-top: 1px solid var(--border);
    font-family: var(--mono); font-size: 11px;
    color: var(--text-soft); letter-spacing: 1px;
    display: flex; justify-content: space-between; flex-wrap: wrap; gap: 12px;
  }}
  footer a {{ color: var(--text-mid); }}
  footer a:hover {{ color: var(--brass); }}

  @media (max-width: 720px) {{
    .wrap {{ padding: 36px 20px 56px; }}
    h1.page-title {{ font-size: 36px; letter-spacing: -1px; }}
  }}

  {extra_css}
</style>
</head>
<body>
<div class="topbar">
  <a class="tb-l" href="/" title="Back to dashboard">
    <svg viewBox="0 0 64 64" aria-hidden="true">
      <path fill-rule="evenodd" clip-rule="evenodd" d="M32 3 L42 22 L61 32 L42 42 L32 61 L22 42 L3 32 L22 22 Z M32 26 A6 6 0 1 0 32 38 A6 6 0 1 0 32 26 Z" fill="currentColor"/>
      <circle cx="32" cy="32" r="2.2" fill="currentColor"/>
    </svg>
    <strong>FORTIFY THE ORDNANCE</strong>
  </a>
  <div class="tb-r"><a href="/">← back to dashboard</a></div>
</div>
<div class="wrap">
{body}
</div>
<footer>
  <span>{footer_left}</span>
  <a href="https://github.com/Smokeybear10/104-RSCH.BOFARCHIVES" target="_blank" rel="noopener">github ↗</a>
</footer>
</body>
</html>"""


def render(title: str, body: str, extra_css: str = "", footer_left: str = "") -> str:
    return _SHELL.format(
        title=h(title), body=body, extra_css=extra_css,
        footer_left=h(footer_left) if footer_left else "Built by Thomas Ou",
    )


# ── 1. CHIEFS OF ORDNANCE (Crozier + Lewis) ─────────────────────────────────

_DATE_RE = re.compile(r"\b(18\d{2}|19[0-3]\d)\b")
_PERIOD_RE = re.compile(r"^\s*(\d{4})(?:[-–]\s*(\d{4}))?\s*[:.\-]?\s*(.*)$")


def _parse_biography(path: Path) -> list[dict]:
    """Walk a Crozier/Lewis docx, group paragraphs into year-keyed events.

    Returns list of {year, headline, detail_lines}.
    """
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Skip header lines about research team, assignment, resources
    start = 0
    for i, p in enumerate(paras):
        if _DATE_RE.search(p) and ("18" in p or "19" in p):
            start = i; break

    events: list[dict] = []
    current: dict | None = None
    for line in paras[start:]:
        # Skip URLs, research instructions
        if line.startswith(("http", "Proquest", "Cullum", "https://")):
            continue
        if line.lower().startswith(("research", "assignment", "initial research", "dates on", "dates/details")):
            continue
        # Period heading like "1876-1877" or "1892-1902: Board of Ordnance (Ellie)"
        m = _PERIOD_RE.match(line)
        is_period = (
            m is not None
            and len(line) < 80
            and m.group(1)
            and not any(month in line.lower() for month in
                        ["january","february","march","april","may","june","july",
                         "august","september","october","november","december",
                         "jan ","feb ","mar ","apr ","jun ","jul ","aug ","sep ","oct ","nov ","dec "])
        )
        if is_period:
            if current and current["detail"]:
                events.append(current)
            year = int(m.group(1))
            label = (m.group(3) or "").strip(" -–:")
            current = {"year": year, "headline": label, "detail": []}
            continue
        # Dated line like "June 15, 1884 – ..." or "1892 – Shift into fire-control"
        if _DATE_RE.search(line):
            # Pull the 4-digit year out
            year_m = _DATE_RE.search(line)
            year = int(year_m.group(1))
            if current and current["detail"]:
                events.append(current)
            current = {"year": year, "headline": line, "detail": []}
            continue
        # Otherwise treat as detail for the current event
        if current is not None:
            current["detail"].append(line)
        else:
            current = {"year": None, "headline": line, "detail": []}
    if current and (current["detail"] or current.get("headline")):
        events.append(current)
    return events


_CHIEFS_CSS = """
  .chief-tabs {
    display: flex; gap: 4px; margin: 32px 0 20px;
    border-bottom: 1px solid var(--border);
  }
  .chief-tab {
    padding: 12px 22px;
    font-family: var(--mono);
    font-size: 11px; letter-spacing: 1.3px;
    text-transform: uppercase; font-weight: 700;
    background: transparent; border: 0; cursor: pointer;
    color: var(--text-mid);
    border-bottom: 2px solid transparent;
    margin-bottom: -1px;
    transition: color .15s, border-color .15s;
  }
  .chief-tab.active { color: var(--text); border-bottom-color: var(--brass); }
  .chief-tab:hover { color: var(--text); }
  .chief-panel { display: none; }
  .chief-panel.active { display: block; }

  .chief-card {
    display: grid;
    grid-template-columns: 220px 1fr;
    gap: 32px;
    padding: 24px 0;
    border-top: 1px solid var(--border);
  }
  .chief-card:first-of-type { border-top: 0; }
  .chief-avatar {
    width: 220px; height: 220px;
    border-radius: 8px;
    background: linear-gradient(135deg, #1D3461 0%, #2D5A8F 100%);
    display: flex; align-items: center; justify-content: center;
    color: white;
    font-family: var(--serif);
    font-size: 64px;
    font-weight: 700;
  }
  .chief-avatar.lewis { background: linear-gradient(135deg, #5D4037 0%, #8C6E5A 100%); }
  .chief-bio h2 {
    font-family: var(--serif);
    font-size: 32px; font-weight: 700; letter-spacing: -0.6px;
    margin-bottom: 4px;
  }
  .chief-bio .rank {
    font-family: var(--mono); font-size: 11px;
    letter-spacing: 1.3px; text-transform: uppercase;
    color: var(--brass); font-weight: 700; margin-bottom: 14px;
  }
  .chief-bio p { color: var(--text-mid); margin-bottom: 10px; }
  .chief-bio .stats { display: flex; gap: 18px; margin-top: 12px; flex-wrap: wrap; }
  .chief-bio .stat {
    background: var(--panel); border: 1px solid var(--border);
    padding: 10px 16px; border-radius: 6px; min-width: 110px;
  }
  .chief-bio .stat .n {
    font-family: var(--serif); font-size: 20px; font-weight: 700;
  }
  .chief-bio .stat .l {
    font-family: var(--mono); font-size: 9.5px;
    letter-spacing: 1.4px; text-transform: uppercase;
    color: var(--text-soft); margin-top: 4px; font-weight: 600;
  }

  .timeline {
    margin-top: 36px;
    padding-left: 0;
    position: relative;
  }
  .timeline::before {
    content: ""; position: absolute;
    top: 0; bottom: 0; left: 110px;
    width: 2px; background: var(--border);
  }
  .event {
    display: grid;
    grid-template-columns: 110px 1fr;
    gap: 28px;
    position: relative;
    padding: 16px 0;
  }
  .event::before {
    content: ""; position: absolute;
    left: 105px; top: 24px;
    width: 12px; height: 12px;
    border-radius: 50%;
    background: var(--brass);
    border: 3px solid var(--bg);
    z-index: 1;
  }
  .event .year {
    font-family: var(--mono);
    font-size: 13px; font-weight: 700; letter-spacing: 0.6px;
    color: var(--text);
    text-align: right;
    padding-top: 18px;
  }
  .event .body {
    background: var(--panel);
    border: 1px solid var(--border);
    border-radius: 6px;
    padding: 14px 18px;
    transition: border-color .15s, box-shadow .15s;
  }
  .event:hover .body {
    border-color: var(--brass);
    box-shadow: 0 4px 12px rgba(60, 40, 10, .08);
  }
  .event .headline {
    font-family: var(--serif);
    font-size: 15.5px; font-weight: 700; line-height: 1.35;
    color: var(--text); margin-bottom: 4px;
  }
  .event .detail {
    font-size: 14px; line-height: 1.55;
    color: var(--text-mid);
  }
  .event .detail p { margin-bottom: 4px; }
  @media (max-width: 720px) {
    .chief-card { grid-template-columns: 1fr; }
    .chief-avatar { width: 100%; height: 160px; }
    .timeline::before { left: 60px; }
    .event { grid-template-columns: 60px 1fr; gap: 14px; }
    .event::before { left: 55px; }
    .event .year { font-size: 11px; }
  }
"""


def build_chiefs() -> None:
    crozier = _parse_biography(DATA / "Crozier and Lewis/Crozier Timeline with details.docx")
    lewis_a = _parse_biography(DATA / "Crozier and Lewis/Lewis Timeline with Details.docx")
    lewis_b = _parse_biography(DATA / "Crozier and Lewis/Ellie Lewis Timeline with Details.docx")
    # Merge the two Lewis sources by year ordering, dedupe-ish (keep both if year differs by a year)
    lewis = sorted(lewis_a + lewis_b, key=lambda e: (e.get("year") or 9999))

    def render_events(events: list[dict]) -> str:
        out = ['<div class="timeline">']
        for e in events:
            year = e.get("year")
            if year is None or year < 1850 or year > 1960:
                continue
            head = e["headline"]
            # Strip leading "YYYY " or "Month dd, YYYY " from headline if year already in side column
            head_clean = re.sub(r"^\s*(\d{4}|\w+\s+\d+,?\s*\d{4}|\w+\s+\d{1,2}(?:st|nd|rd|th)?)\s*[-–:]?\s*", "", head).strip()
            if not head_clean:
                head_clean = head
            detail_html = "".join(f"<p>{h(d)}</p>" for d in e["detail"][:6])
            out.append(
                f'<div class="event">'
                f'<div class="year">{year}</div>'
                f'<div class="body">'
                f'<div class="headline">{h(head_clean[:200])}</div>'
                + (f'<div class="detail">{detail_html}</div>' if detail_html else '')
                + '</div></div>'
            )
        out.append('</div>')
        return "\n".join(out)

    cr_events = [e for e in crozier if e.get("year") and 1850 <= e["year"] <= 1960]
    le_events = [e for e in lewis   if e.get("year") and 1850 <= e["year"] <= 1960]
    cr_years = [e["year"] for e in cr_events]
    le_years = [e["year"] for e in le_events]

    body = f'''
  <div class="eyebrow">Chiefs of Ordnance</div>
  <h1 class="page-title">The two men who <em>ran</em> the Board</h1>
  <p class="lede">
    Brig. Gen. William Crozier sat as Chief of Ordnance for 17 years (1901–1918), spanning the
    Spanish-American War, the run-up to WWI, and the U.S. mobilization itself.
    Colonel Isaac Newton Lewis was the Board's recorder and the inventor whose work crystallized
    coast-artillery fire control. Their careers, reconstructed from RA timelines.
  </p>

  <div class="chief-tabs" role="tablist">
    <button class="chief-tab active" data-chief="crozier" role="tab">William Crozier</button>
    <button class="chief-tab" data-chief="lewis" role="tab">Isaac N. Lewis</button>
  </div>

  <div class="chief-panel active" id="panel-crozier">
    <div class="chief-card">
      <div class="chief-avatar">WC</div>
      <div class="chief-bio">
        <h2>William Crozier</h2>
        <div class="rank">Brig. Gen. · Chief of Ordnance 1901–1918</div>
        <p>USMA 1876 (5th in class). Joint inventor of the Buffington–Crozier disappearing carriage,
           designer of most of the service siege and seacoast carriages, and inventor of the Crozier
           wire-wound gun. Served as U.S. delegate to the Hague Peace Conference, Chief of Ordnance
           through both presidencies of Theodore Roosevelt and William Howard Taft, and led ordnance
           procurement during America's WWI mobilization.</p>
        <div class="stats">
          <div class="stat"><div class="n">{len(cr_events)}</div><div class="l">Career events</div></div>
          <div class="stat"><div class="n">{min(cr_years) if cr_years else '—'}–{max(cr_years) if cr_years else '—'}</div><div class="l">Span</div></div>
        </div>
      </div>
    </div>
    {render_events(cr_events)}
  </div>

  <div class="chief-panel" id="panel-lewis">
    <div class="chief-card">
      <div class="chief-avatar lewis">IL</div>
      <div class="chief-bio">
        <h2>Isaac Newton Lewis</h2>
        <div class="rank">Col. · BOF recorder · Inventor</div>
        <p>USMA 1884 (graduating with Crozier's contemporaries). Invented the Lewis range and
           position finder while stationed at Fort Leavenworth (1888–1890), formally adopted by the
           War Department in 1896 as the basic instrument for U.S. coast artillery fire control.
           Designer of the Lewis Sub-Scale verniers and an early plotting board for battery fire
           control. Later inventor of the Lewis automatic machine gun, which he eventually licensed
           abroad after the Board declined to adopt it.</p>
        <div class="stats">
          <div class="stat"><div class="n">{len(le_events)}</div><div class="l">Career events</div></div>
          <div class="stat"><div class="n">{min(le_years) if le_years else '—'}–{max(le_years) if le_years else '—'}</div><div class="l">Span</div></div>
        </div>
      </div>
    </div>
    {render_events(le_events)}
  </div>
'''

    js = """
<script>
  document.querySelectorAll('.chief-tab').forEach(tab => {
    tab.addEventListener('click', () => {
      const key = tab.dataset.chief;
      document.querySelectorAll('.chief-tab').forEach(t => t.classList.toggle('active', t === tab));
      document.querySelectorAll('.chief-panel').forEach(p => p.classList.toggle('active', p.id === 'panel-' + key));
      window.scrollTo({top: 0, behavior: 'smooth'});
    });
  });
</script>
"""

    page = render("Chiefs of Ordnance", body + js, _CHIEFS_CSS,
                  f"Parsed from Data/Crozier and Lewis/ · {len(cr_events)} Crozier events · {len(le_events)} Lewis events")
    (ROOT / "chiefs.html").write_text(page, encoding="utf-8")
    print(f"  ✓ chiefs.html ({len(cr_events)} Crozier + {len(le_events)} Lewis events)")


# ── 2. REPORTS — Annual Report excerpts ─────────────────────────────────────

def _parse_reports(path: Path) -> list[dict]:
    """Walk Non-Gun docx, group into Annual Report sections × topic excerpts."""
    doc = Document(path)
    reports: list[dict] = []
    current_ar: dict | None = None
    current_topic: dict | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        style = para.style.name

        # Annual Report header
        m = re.match(r"^Annual\s+Report\s+(\d+)\s*\(?(\d{4})\s*-\s*(\d{4})\)?", text)
        if m and (style.startswith("Heading") or style == "Title"):
            if current_topic and current_topic["body"]:
                current_ar["topics"].append(current_topic)
            if current_ar:
                reports.append(current_ar)
            current_ar = {
                "number": int(m.group(1)),
                "year_start": int(m.group(2)),
                "year_end": int(m.group(3)),
                "topics": [],
            }
            current_topic = None
            continue

        # Topic header — "Topic Name (pg. NN)"
        tm = re.match(r"^(.+?)\s*\(pg\.?\s*([\dvixIVX]+)\)\s*$", text, re.IGNORECASE)
        if tm and current_ar:
            if current_topic and current_topic["body"]:
                current_ar["topics"].append(current_topic)
            current_topic = {
                "name": tm.group(1).strip(),
                "page": tm.group(2).strip(),
                "body": [],
            }
            continue

        # "Appropriations & Funding" / "BOF Excerpts" subsection markers
        if text in {"Appropriations & Funding", "BOF Excerpts"} and current_ar:
            continue

        # Otherwise body text
        if current_topic is not None:
            current_topic["body"].append(text)
        elif current_ar is not None:
            # Body before any topic — make a default topic
            current_topic = {"name": "Overview", "page": "", "body": [text]}

    if current_topic and current_topic["body"]:
        current_ar["topics"].append(current_topic)
    if current_ar:
        reports.append(current_ar)
    return reports


_REPORTS_CSS = """
  .ar-grid {
    display: grid;
    grid-template-columns: 280px 1fr;
    gap: 48px;
    margin-top: 36px;
  }
  .ar-nav {
    position: sticky; top: 24px;
    align-self: start;
    max-height: calc(100vh - 48px);
    overflow-y: auto;
  }
  .ar-nav button {
    display: block; width: 100%;
    padding: 12px 14px;
    background: transparent; border: 0; cursor: pointer;
    text-align: left;
    font-family: var(--sans);
    font-size: 13.5px; line-height: 1.35;
    color: var(--text-mid);
    border-left: 3px solid transparent;
    transition: background .12s, color .12s, border-color .12s;
  }
  .ar-nav button:hover { background: var(--panel); color: var(--text); }
  .ar-nav button.active {
    color: var(--text); border-left-color: var(--brass);
    background: var(--panel);
    font-weight: 600;
  }
  .ar-nav button .num {
    font-family: var(--mono); font-size: 10.5px;
    color: var(--brass); letter-spacing: 1px;
    font-weight: 700;
    display: block; margin-bottom: 2px;
  }

  .ar-panel { display: none; }
  .ar-panel.active { display: block; }
  .ar-header {
    margin-bottom: 18px;
    padding-bottom: 16px;
    border-bottom: 1px solid var(--border);
  }
  .ar-header h2 {
    font-family: var(--serif);
    font-size: 36px; font-weight: 700; letter-spacing: -0.8px;
  }
  .ar-header .range {
    font-family: var(--mono);
    font-size: 11px; letter-spacing: 1.4px;
    text-transform: uppercase; color: var(--brass);
    font-weight: 700; margin-top: 6px;
  }
  .topic {
    margin-bottom: 28px;
    padding: 22px 26px;
    background: var(--panel);
    border: 1px solid var(--border);
    border-radius: 8px;
    border-left: 4px solid var(--brass);
  }
  .topic h3 {
    font-family: var(--serif);
    font-size: 20px; font-weight: 700;
    margin-bottom: 4px;
  }
  .topic .page-ref {
    font-family: var(--mono); font-size: 10.5px;
    letter-spacing: 1px; text-transform: uppercase;
    color: var(--text-soft); margin-bottom: 14px;
    font-weight: 600;
  }
  .topic p {
    font-family: var(--serif);
    font-size: 16px; line-height: 1.65;
    color: var(--text); margin-bottom: 12px;
  }
  .topic p:last-child { margin-bottom: 0; }

  @media (max-width: 980px) {
    .ar-grid { grid-template-columns: 1fr; }
    .ar-nav { position: relative; max-height: none; }
  }
"""


def build_reports() -> None:
    reports = _parse_reports(DATA / "Misc/BOF Non-Gun Assignment_complete.docx")
    # Filter out empty reports (no topics with body)
    reports = [r for r in reports if any(t["body"] for t in r["topics"])]

    nav = '<div class="ar-nav">'
    for i, r in enumerate(reports):
        cls = "active" if i == 0 else ""
        nav += (
            f'<button class="{cls}" data-ar="{r["number"]}">'
            f'<span class="num">AR {r["number"]}</span>'
            f'{r["year_start"]}–{r["year_end"]}'
            f'</button>'
        )
    nav += '</div>'

    panels = ""
    for i, r in enumerate(reports):
        cls = "active" if i == 0 else ""
        topics_html = ""
        for t in r["topics"]:
            if not t["body"]: continue
            body_html = "".join(f"<p>{h(b)}</p>" for b in t["body"])
            page_ref = f"Page {t['page']}" if t["page"] else "Annual Report"
            topics_html += (
                f'<div class="topic">'
                f'<h3>{h(t["name"])}</h3>'
                f'<div class="page-ref">{h(page_ref)}</div>'
                f'{body_html}'
                f'</div>'
            )
        panels += (
            f'<div class="ar-panel {cls}" id="ar-{r["number"]}">'
            f'<div class="ar-header">'
            f'<h2>Annual Report {r["number"]}</h2>'
            f'<div class="range">{r["year_start"]}–{r["year_end"]} · {len([t for t in r["topics"] if t["body"]])} topics</div>'
            f'</div>'
            f'{topics_html}'
            f'</div>'
        )

    body = f'''
  <div class="eyebrow">Source texts · Annual Reports</div>
  <h1 class="page-title">The Board <em>in</em> its own words</h1>
  <p class="lede">
    Direct excerpts from the BOF Annual Reports — the prose the Board itself wrote about
    submarine mines, range finders, smokeless powders, telegraphic communications, and every
    other subject it deliberated on. Click a report on the left to read.
  </p>
  <div class="ar-grid">
    {nav}
    <div class="ar-content">{panels}</div>
  </div>
'''

    js = """
<script>
  document.querySelectorAll('.ar-nav button').forEach(btn => {
    btn.addEventListener('click', () => {
      const key = btn.dataset.ar;
      document.querySelectorAll('.ar-nav button').forEach(b => b.classList.toggle('active', b === btn));
      document.querySelectorAll('.ar-panel').forEach(p => p.classList.toggle('active', p.id === 'ar-' + key));
      window.scrollTo({top: 0, behavior: 'smooth'});
    });
  });
</script>
"""

    n_topics = sum(len([t for t in r["topics"] if t["body"]]) for r in reports)
    page = render("BOF Source Texts", body + js, _REPORTS_CSS,
                  f"Parsed from Data/Misc/BOF Non-Gun Assignment_complete.docx · {len(reports)} reports · {n_topics} topics")
    (ROOT / "reports.html").write_text(page, encoding="utf-8")
    print(f"  ✓ reports.html ({len(reports)} reports, {n_topics} topics)")


# ── 3. TECH MASTER — year-period × tech mapping ─────────────────────────────

def _parse_tech_mapping(path: Path) -> list[dict]:
    """Walk Technology Mapping docx, group entries by year-period heading."""
    doc = Document(path)
    periods: list[dict] = []
    current: dict | None = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        style = para.style.name

        # Period heading like "1888-1891" (typically Heading 3)
        m = re.match(r"^\s*(\d{4})\s*[-–]\s*(\d{4})\s*$", text)
        if m and (style.startswith("Heading") or style == "Title"):
            if current and current["techs"]:
                periods.append(current)
            current = {
                "label": f"{m.group(1)}–{m.group(2)}",
                "start": int(m.group(1)),
                "end": int(m.group(2)),
                "techs": [],
            }
            continue

        # Skip the leading instructions section
        if current is None:
            continue
        # Skip obvious instruction lines
        if any(t in text.lower() for t in (
            "color coding", "big case study", "everything crozier",
            "substantive flair", "something fun maybe",
        )):
            continue
        # Otherwise it's a tech entry
        current["techs"].append(text)

    if current and current["techs"]:
        periods.append(current)
    return periods


_TECH_CSS = """
  .tech-grid {
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(360px, 1fr));
    gap: 18px;
    margin-top: 36px;
  }
  .tech-period {
    background: var(--panel);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 22px 24px;
    transition: border-color .15s, box-shadow .15s;
  }
  .tech-period:hover {
    border-color: var(--brass);
    box-shadow: 0 4px 14px rgba(60, 40, 10, .08);
  }
  .tech-period .label {
    font-family: var(--mono);
    font-size: 11px; font-weight: 700; letter-spacing: 1.4px;
    text-transform: uppercase; color: var(--brass);
    margin-bottom: 4px;
  }
  .tech-period h3 {
    font-family: var(--serif);
    font-size: 22px; font-weight: 700; letter-spacing: -0.4px;
    margin-bottom: 14px;
  }
  .tech-period .count {
    font-family: var(--mono);
    font-size: 10.5px; letter-spacing: 0.8px;
    color: var(--text-soft); font-weight: 600;
    text-transform: uppercase;
    margin-bottom: 12px;
  }
  .tech-period ul {
    list-style: none;
    padding: 0; margin: 0;
  }
  .tech-period li {
    padding: 6px 0;
    font-size: 14px;
    color: var(--text);
    border-bottom: 1px solid var(--border);
    line-height: 1.4;
  }
  .tech-period li:last-child { border-bottom: 0; }
  .tech-period li::before {
    content: "▸";
    color: var(--brass);
    margin-right: 8px;
    font-size: 10px;
  }
"""


def build_tech_master() -> None:
    periods = _parse_tech_mapping(DATA / "Technologies/Technology Mapping (BOF Annual Reports).docx")
    total_techs = sum(len(p["techs"]) for p in periods)

    cards = ""
    for p in periods:
        techs_html = "".join(f"<li>{h(t[:140])}</li>" for t in p["techs"][:40])
        if len(p["techs"]) > 40:
            techs_html += f'<li style="color:var(--text-soft);font-style:italic;">…and {len(p["techs"])-40} more</li>'
        cards += (
            f'<div class="tech-period">'
            f'<div class="label">{h(p["label"])}</div>'
            f'<h3>BOF AR {p["label"]}</h3>'
            f'<div class="count">{len(p["techs"])} subjects mapped</div>'
            f'<ul>{techs_html}</ul>'
            f'</div>'
        )

    body = f'''
  <div class="eyebrow">Technology mapping</div>
  <h1 class="page-title">The full <em>roster</em></h1>
  <p class="lede">
    The complete list of technologies the BOF reviewed, mapped to the annual report period
    in which each was discussed. Pulled from the canonical assignment doc — broader than the
    subjects-considered XLSXs (which only cover 1897–1908) because it extends back through
    every BOF annual report from 1888.
  </p>
  <p style="font-family: var(--mono); font-size: 11.5px; color: var(--text-soft); letter-spacing: 0.4px;">
    <strong style="color: var(--text-mid);">{len(periods)} report periods</strong>
    &nbsp;·&nbsp;
    <strong style="color: var(--text-mid);">{total_techs} total subject entries</strong>
  </p>
  <div class="tech-grid">{cards}</div>
'''
    page = render("Technology Mapping", body, _TECH_CSS,
                  f"Parsed from Data/Technologies/Technology Mapping (BOF Annual Reports).docx · {total_techs} entries")
    (ROOT / "tech-master.html").write_text(page, encoding="utf-8")
    print(f"  ✓ tech-master.html ({len(periods)} periods, {total_techs} entries)")


# ── 4. READING ROOM — index page linking all parsed docs ─────────────────────

_READING_CSS = """
  .reading-grid {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(380px, 1fr));
    gap: 18px;
    margin-top: 36px;
  }
  .reading-card {
    background: var(--panel);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 28px 30px;
    text-decoration: none;
    color: inherit;
    transition: border-color .15s, box-shadow .15s, transform .12s;
    display: block;
  }
  .reading-card:hover {
    border-color: var(--brass);
    box-shadow: 0 6px 20px rgba(60, 40, 10, .10);
    transform: translateY(-2px);
    text-decoration: none;
  }
  .reading-card .tag {
    font-family: var(--mono);
    font-size: 10.5px; font-weight: 700; letter-spacing: 1.4px;
    text-transform: uppercase; color: var(--brass);
    margin-bottom: 12px;
  }
  .reading-card h2 {
    font-family: var(--serif);
    font-size: 26px; font-weight: 700; letter-spacing: -0.5px;
    line-height: 1.15; margin-bottom: 8px;
  }
  .reading-card p {
    color: var(--text-mid);
    font-size: 14.5px; line-height: 1.55;
    margin-bottom: 14px;
  }
  .reading-card .meta {
    font-family: var(--mono);
    font-size: 10.5px; letter-spacing: 0.8px;
    color: var(--text-soft); font-weight: 600;
    text-transform: uppercase;
  }
"""


def build_reading_room() -> None:
    body = '''
  <div class="eyebrow">Reading room</div>
  <h1 class="page-title">The <em>longer</em> records</h1>
  <p class="lede">
    The dashboard summarizes numbers; the reading room is where the actual prose lives.
    Biographies, source texts, technology rosters, RA memos — everything parsed from the
    original Word / PDF files in the research drive.
  </p>
  <div class="reading-grid">

    <a class="reading-card" href="/chiefs.html">
      <div class="tag">People</div>
      <h2>Chiefs of Ordnance</h2>
      <p>Annotated career timelines for Brig. Gen. William Crozier and Col. Isaac N. Lewis — the
         two figures who shaped the Board for two decades. Side-by-side tabs, year-anchored events.</p>
      <div class="meta">Crozier and Lewis docs · biographical</div>
    </a>

    <a class="reading-card" href="/reports.html">
      <div class="tag">Primary source</div>
      <h2>BOF Source Texts</h2>
      <p>Direct excerpts from each Annual Report — submarine mines, fire-control systems, smokeless
         powders, fuses, telegraphic communication. The Board's prose at the moment of each decision.</p>
      <div class="meta">Non-Gun Assignment doc · 8 annual reports</div>
    </a>

    <a class="reading-card" href="/tech-master.html">
      <div class="tag">Reference</div>
      <h2>Technology Mapping</h2>
      <p>Master roster of every technology mapped to its BOF Annual Report period. Broader
         than the proposals corpus — extends back to 1888 instead of starting at 1897.</p>
      <div class="meta">Technology Mapping doc · canonical roster</div>
    </a>

    <a class="reading-card" href="/tech-memos.html">
      <div class="tag">RA write-ups</div>
      <h2>Technology Memos</h2>
      <p>Detailed RA write-ups on specific technologies — Howell torpedoes, Luger pistol,
         Lewis range finder, Brown segmental gun, anti-submarine net. Each card opens a full memo.</p>
      <div class="meta">Technology Memos folder · 6 filled, more in progress</div>
    </a>

    <a class="reading-card" href="/data-quality.html">
      <div class="tag">Honest accounting</div>
      <h2>Data Quality</h2>
      <p>Every gap and quirk in the underlying data, listed with file paths. What's missing,
         what's transcription noise, and what was already fixed.</p>
      <div class="meta">Audited 2026-05-12</div>
    </a>

    <a class="reading-card" href="/about.html">
      <div class="tag">Project</div>
      <h2>About</h2>
      <p>What the Board of Ordnance & Fortification was, who built this dashboard, and where
         the source material came from.</p>
      <div class="meta">Background · team · methodology</div>
    </a>

  </div>
'''
    page = render("Reading Room", body, _READING_CSS,
                  "Index · long-form content from the research drive")
    (ROOT / "reading-room.html").write_text(page, encoding="utf-8")
    print(f"  ✓ reading-room.html")


# ── Main ────────────────────────────────────────────────────────────────────

def main() -> None:
    print("Parsing the research drive…\n")
    build_chiefs()
    build_reports()
    build_tech_master()
    build_reading_room()
    print("\nDone. Pages: chiefs.html, reports.html, tech-master.html, reading-room.html")


if __name__ == "__main__":
    main()
